import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

CHAPTERS_DIR = r"c:\Users\tsuba\GolandProjects\InconsistencyFixer\output\writer\Completed-stories\vow-of-the-bloodthorne\chapters-revised"
OUTPUT_PATH = r"c:\Users\tsuba\GolandProjects\InconsistencyFixer\output\writer\Completed-stories\vow-of-the-bloodthorne\Vow_of_the_Bloodthorne.docx"

def get_sorted_chapters(directory):
    files = []
    for fname in os.listdir(directory):
        if fname.endswith(".txt") and "DUPLICATE" not in fname:
            match = re.match(r"chapter_(\d+)\.txt", fname)
            if match:
                files.append((int(match.group(1)), os.path.join(directory, fname)))
    files.sort(key=lambda x: x[0])
    return files

def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = Pt(18)
    style.paragraph_format.space_after = Pt(0)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

def add_chapter(doc, num, filepath):
    with open(filepath, encoding="utf-8") as f:
        content = f.read().strip()

    lines = content.split("\n")
    first_line = lines[0].strip() if lines else ""

    # Chapter heading
    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.runs[0] if heading.runs else heading.add_run(first_line)
    if not heading.runs:
        pass
    else:
        run.text = first_line
    heading.style.font.name = "Georgia"

    # Body paragraphs
    body_lines = lines[1:]
    i = 0
    while i < len(body_lines):
        line = body_lines[i].strip()
        if not line:
            i += 1
            continue
        # Scene break
        if line in ("* * *", "***", "---"):
            p = doc.add_paragraph("* * *")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Pt(36)
        i += 1

    # Page break between chapters
    doc.add_page_break()

def main():
    chapters = get_sorted_chapters(CHAPTERS_DIR)
    print(f"Found {len(chapters)} chapters.")

    doc = Document()
    setup_styles(doc)

    # Title page
    title = doc.add_heading("Vow of the Bloodthorne", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for num, path in chapters:
        print(f"  Adding chapter {num:04d}...")
        add_chapter(doc, num, path)

    doc.save(OUTPUT_PATH)
    print(f"\nSaved: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
